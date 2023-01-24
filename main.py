import docx

# Abrir el archivo de Word
doc = docx.Document('prueba.docx')

# Recorrer todas las oraciones del documento
for para in doc.paragraphs:
    # Recorrer todas las palabras de cada oración
    for run in para.runs:
        # Reemplazar el símbolo * con una imagen
        if '*' in run.text:
            run.text = run.text.replace('*', '')
            run.add_picture('image.jpg')

# Guardar el archivo de Word con las modificaciones
doc.save('modified_example.docx')